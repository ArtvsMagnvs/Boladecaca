# tests/test_lectura_paginada.py — documentos largos (2026-08-02)
#
# EL FALLO QUE CIERRA (reportado varias veces por el usuario, la última con el
# GDD de Cordyceps): "el contenido del documento se cortó durante la lectura".
# Era cierto y no tenía salida — la tool devolvía el texto entero, el toolloop
# lo recortaba con un aviso honesto, y el modelo no tenía forma de pedir el
# resto. Ahora la lectura es paginada y el aviso dice cómo continuar.
from __future__ import annotations

import asyncio
import os

import pytest

from app.tie import toolloop
from app.tools.document_tool import (DEFAULT_READ_CHARS, MAX_READ_CHARS,
                                     DocumentTool, _window)
from app.tools.filesystem_tool import FilesystemTool


# ===========================================================================
# 1 — La ventana: función pura, sin perder ni un carácter
# ===========================================================================
class TestVentana:
    def test_texto_corto_cabe_de_una_vez(self):
        w = _window("hola", {})
        assert w["text"] == "hola" and w["has_more"] is False
        assert w["next_offset"] is None and w["total_chars"] == 4

    def test_recorrer_un_texto_largo_lo_recupera_entero(self):
        texto = "\n".join(f"linea {i} con contenido de relleno" for i in range(3000))
        partes, off, vueltas = [], 0, 0
        while True:
            w = _window(texto, {"offset": off})
            partes.append(w["text"])
            vueltas += 1
            assert vueltas < 100, "bucle infinito"
            if not w["has_more"]:
                break
            off = w["next_offset"]
        assert vueltas > 1, "el texto de prueba debería necesitar varias ventanas"
        recuperado = "".join(partes)
        # Se pierden solo los '\n' de los cortes (uno por ventana como mucho).
        assert len(recuperado) >= len(texto) - vueltas
        for marca in ("linea 0 ", "linea 1500 ", "linea 2999 "):
            assert marca in recuperado

    def test_el_corte_respeta_los_saltos_de_linea(self):
        """Ninguna línea puede quedar partida entre dos ventanas: al recorrer
        el texto entero, todas las líneas originales deben aparecer completas."""
        lineas = [f"linea-{i}-" + "x" * 90 for i in range(400)]
        texto = "\n".join(lineas)
        partes, off = [], 0
        while True:
            w = _window(texto, {"offset": off, "max_chars": 5000})
            partes.append(w["text"])
            if not w["has_more"]:
                break
            off = w["next_offset"]
        assert len(partes) > 1
        for parte in partes:
            for linea in parte.split("\n"):
                if linea:
                    assert linea in lineas, f"línea partida a la mitad: {linea[:40]!r}"

    def test_offset_mas_alla_del_final_no_revienta(self):
        w = _window("corto", {"offset": 9999})
        assert w["text"] == "" and w["has_more"] is False

    @pytest.mark.parametrize("valor", ["basura", None, -5, 10 ** 9])
    def test_parametros_invalidos_degradan_a_algo_sensato(self, valor):
        w = _window("abc" * 1000, {"offset": valor, "max_chars": valor})
        assert isinstance(w["text"], str)
        assert 0 <= w["offset"] <= w["total_chars"]

    def test_max_chars_tiene_techo(self):
        texto = "a" * (MAX_READ_CHARS * 3)
        w = _window(texto, {"max_chars": MAX_READ_CHARS * 2})
        assert len(w["text"]) <= MAX_READ_CHARS


# ===========================================================================
# 2 — Las tools reales: un .docx largo se lee ENTERO en varias llamadas
# ===========================================================================
@pytest.mark.anyio
async def test_un_docx_largo_se_lee_entero_por_partes(tmp_path, monkeypatch):
    monkeypatch.setenv("HOME", str(tmp_path))
    monkeypatch.setattr("app.tools.filesystem_tool.Path.home", lambda: tmp_path, raising=False)
    import docx

    ruta = tmp_path / "GDD.docx"
    d = docx.Document()
    for i in range(1200):
        d.add_paragraph(f"Parrafo {i}: Cordyceps es un roguelike 2D con estetica pixel art.")
    d.save(str(ruta))

    tool = DocumentTool()
    leido, off, llamadas = "", 0, 0
    while True:
        r = await tool.execute("read_docx", {"path": str(ruta), "offset": off})
        assert r["success"], r
        res = r["result"]
        leido += res["text"]
        llamadas += 1
        assert llamadas < 40
        if not res["has_more"]:
            break
        off = res["next_offset"]

    assert llamadas > 1, "el documento de prueba debería necesitar varias lecturas"
    assert "Parrafo 0:" in leido and "Parrafo 1199:" in leido, "no llegó al final"


@pytest.mark.anyio
async def test_la_primera_ventana_avisa_de_como_seguir(tmp_path, monkeypatch):
    monkeypatch.setenv("HOME", str(tmp_path))
    import docx

    ruta = tmp_path / "largo.docx"
    d = docx.Document()
    for i in range(1200):
        d.add_paragraph(f"Parrafo {i} con bastante texto de relleno para pasar del limite.")
    d.save(str(ruta))

    res = (await DocumentTool().execute("read_docx", {"path": str(ruta)}))["result"]
    assert res["has_more"] is True
    assert "offset=" in res["note"], "la nota no dice cómo continuar"
    assert str(res["next_offset"]) in res["note"]


@pytest.mark.anyio
async def test_read_file_pagina_y_sigue_siendo_compatible(tmp_path, monkeypatch):
    monkeypatch.setenv("HOME", str(tmp_path))
    ruta = tmp_path / "grande.txt"
    ruta.write_text("\n".join(f"fila {i}" for i in range(8000)), encoding="utf-8")

    fs = FilesystemTool()
    # Sin parámetros nuevos: el contrato de siempre sigue funcionando.
    r = await fs.execute("read_file", {"path": str(ruta)})
    assert r["success"] and r["result"]["content"]
    assert r["result"]["has_more"] is True

    # Con offset: continúa donde lo dejó.
    r2 = await fs.execute("read_file", {"path": str(ruta), "offset": r["result"]["next_offset"]})
    assert r2["result"]["offset"] == r["result"]["next_offset"]
    assert r2["result"]["content"] not in ("", r["result"]["content"])


# ===========================================================================
# 3 — La observación del toolloop: el aviso es ACCIONABLE
# ===========================================================================
class TestObservacionAccionable:
    def test_dice_la_llamada_exacta_para_continuar(self):
        payload = {"path": "/x/GDD.docx", "text": "contenido leido",
                   "has_more": True, "next_offset": 20000, "total_chars": 91234}
        obs = toolloop._observation("document", "read_docx", payload)
        assert "offset=20000" in obs
        assert "read_docx" in obs
        assert "NO respondas todavía" in obs

    def test_sin_resto_no_mete_ruido(self):
        payload = {"path": "/x/corto.docx", "text": "todo el contenido",
                   "has_more": False, "next_offset": None, "total_chars": 17}
        obs = toolloop._observation("document", "read_docx", payload)
        assert "offset=" not in obs and "CONTINUACIÓN" not in obs

    def test_el_contenido_no_viaja_como_json_descabezado(self):
        """[S5] El texto se entrega en plano; el JSON solo era ruido que se
        comía el presupuesto."""
        payload = {"path": "/x/a.txt", "text": "PALABRA_CLAVE " * 100,
                   "has_more": False, "next_offset": None, "total_chars": 1400}
        obs = toolloop._observation("filesystem", "read_file", payload)
        assert obs.count("PALABRA_CLAVE") > 50
        assert '\\"text\\"' not in obs


# ===========================================================================
# 4 — Presupuesto de vueltas: leer por partes tiene que CABER
# ===========================================================================
def test_document_tiene_presupuesto_para_varias_lecturas():
    from app.core.config import settings
    from app.tie.runtime import _iters_for

    assert _iters_for(["document"]) == settings.TIE_TOOL_MAX_ITERS_WRITE
    assert _iters_for(["document"]) > settings.TIE_TOOL_MAX_ITERS, (
        "con el presupuesto de solo-lectura un documento largo se queda a medias"
    )
    # No-regresión: una tool de solo lectura ligera sigue con el presupuesto corto.
    assert _iters_for(["search"]) == settings.TIE_TOOL_MAX_ITERS
