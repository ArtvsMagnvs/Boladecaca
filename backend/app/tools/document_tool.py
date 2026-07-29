# backend/app/tools/document_tool.py
#
# V1.0 (Tools, #218): documentos de oficina REALES. Antes Aithera no tenia
# ninguna forma de leer un PDF ni de producir un XLSX/DOCX de verdad -- un
# hueco de capacidad basico para un asistente (leer un informe en PDF, entregar
# una tabla en Excel, redactar un documento en Word).
#
# Alcance HONESTO (nombre de la tarea: "leer PDF y escribir XLSX/DOCX"):
#   - PDF: solo LECTURA. pypdf extrae el TEXTO de un PDF de texto. Un PDF
#     escaneado (solo imagen) no tiene texto que extraer -> eso es OCR
#     (desktop_tool/winocr), un camino distinto; se avisa en el resultado.
#     Generar PDFs (reportlab/weasyprint) es una bestia mucho mas pesada y menos
#     necesaria -> fuera de alcance a proposito.
#   - DOCX y XLSX: LECTURA y ESCRITURA (simetrico -- las mismas librerias hacen
#     ambas y si Aithera escribe un DOCX, poder leer uno es lo esperado).
#
# Las 3 librerias (pypdf / python-docx / openpyxl) son Python PURO y se importan
# de forma LAZY (igual que Playwright en browser_tool y pyautogui en
# desktop_tool): el ToolManager registra TODAS las tools al arrancar, asi que un
# import a nivel de modulo cargaria lxml/pypdf en cada arranque aunque nadie use
# documentos.
#
# Seguridad: reusa EXACTAMENTE la validacion de paths de FilesystemTool (solo
# dentro de HOME, sin path traversal, sin zonas del sistema) -- la misma que ya
# usan email_tool (adjuntos) y download_tool. Las acciones de escritura piden
# confirmacion (gobernadas por el permiso `filesystem.write`, igual que
# filesystem.write_file).
#
# Acciones:
#   read_pdf    -> texto por pagina (rango opcional)         [lectura]
#   read_docx   -> parrafos + tablas                          [lectura]
#   read_xlsx   -> filas por hoja                             [lectura]
#   write_docx  -> crea un .docx desde bloques estructurados  [escritura, confirmar]
#   write_xlsx  -> crea un .xlsx desde filas/hojas            [escritura, confirmar]

import asyncio
from typing import Dict, Any, List, Optional

from .base import BaseTool
from .filesystem_tool import _resolve_user_path, _is_path_allowed

# Limites (protegen memoria y el contexto del LLM: volcar un libro de 500
# paginas o un Excel de 1M filas al modelo no ayuda a nadie).
MAX_DOC_BYTES = 25 * 1024 * 1024        # 25 MB por archivo de entrada
MAX_PDF_PAGES = 200                     # paginas leidas de un tiron como maximo
MAX_TEXT_CHARS = 500_000               # tope de texto extraido devuelto
MAX_XLSX_ROWS = 5000                    # filas leidas/escritas por hoja
MAX_XLSX_COLS = 256
MAX_WRITE_BLOCKS = 5000                 # bloques de un write_docx


def _check_readable(path_str: str):
    """Resuelve y valida un path de ENTRADA. Devuelve (path, error_dict|None)."""
    if not path_str:
        return None, {"success": False, "result": None, "error": "falta parametro: path"}
    path = _resolve_user_path(path_str)
    if not _is_path_allowed(path):
        return None, {"success": False, "result": None, "error": f"path fuera de zonas permitidas: {path}"}
    if not path.exists():
        return None, {"success": False, "result": None, "error": f"no existe: {path}"}
    if path.is_dir():
        return None, {"success": False, "result": None, "error": f"es un directorio: {path}"}
    size = path.stat().st_size
    if size > MAX_DOC_BYTES:
        return None, {"success": False, "result": None,
                      "error": f"archivo demasiado grande ({size} bytes, max {MAX_DOC_BYTES})"}
    return path, None


def _check_writable(path_str: str):
    """Resuelve y valida un path de SALIDA. Devuelve (path, error_dict|None)."""
    if not path_str:
        return None, {"success": False, "result": None, "error": "falta parametro: path"}
    path = _resolve_user_path(path_str)
    if not _is_path_allowed(path):
        return None, {"success": False, "result": None, "error": f"path fuera de zonas permitidas: {path}"}
    if path.exists() and path.is_dir():
        return None, {"success": False, "result": None, "error": f"es un directorio: {path}"}
    return path, None


class DocumentTool(BaseTool):
    tool_id = "document"
    name = "Document Tool"
    description = (
        "Lee PDF (texto), lee y escribe DOCX (Word) y XLSX (Excel). "
        "Leer no requiere confirmacion; escribir crea un archivo en HOME y si."
    )
    requires_confirmation = False  # depende de la accion

    async def execute(self, action: str, params: Dict[str, Any]) -> Dict[str, Any]:
        try:
            handler = {
                "read_pdf": self._read_pdf,
                "read_docx": self._read_docx,
                "read_xlsx": self._read_xlsx,
                "write_docx": self._write_docx,
                "write_xlsx": self._write_xlsx,
            }.get(action)
            if not handler:
                return {
                    "success": False, "result": None,
                    "error": (f"Accion desconocida: {action}. Disponibles: "
                              "read_pdf, read_docx, read_xlsx, write_docx, write_xlsx"),
                }
            return await handler(params)
        except ValueError as e:
            return {"success": False, "result": None, "error": f"parametro invalido: {e}"}
        except Exception as e:
            return {"success": False, "result": None, "error": f"{type(e).__name__}: {e}"}

    def list_actions(self) -> List[Dict[str, Any]]:
        return [
            {
                "id": "read_pdf",
                "description": "Extrae el texto de un PDF (de texto; un PDF escaneado no tiene texto que extraer).",
                "requires_confirmation": False,
                "params": {
                    "path": "string (path a un .pdf, absoluto o relativo a HOME)",
                    "pages": "string opcional, rango 1-indexado (ej. '1-5' o '3'); default: todo (max 200 pag)",
                },
            },
            {
                "id": "read_docx",
                "description": "Lee un documento Word: devuelve sus parrafos y el contenido de sus tablas.",
                "requires_confirmation": False,
                "params": {"path": "string (path a un .docx)"},
            },
            {
                "id": "read_xlsx",
                "description": "Lee un Excel: devuelve las filas de cada hoja (o de una hoja concreta).",
                "requires_confirmation": False,
                "params": {
                    "path": "string (path a un .xlsx)",
                    "sheet": "string opcional (nombre de una hoja; default: todas)",
                },
            },
            {
                "id": "write_docx",
                "description": "Crea un documento Word desde bloques estructurados (titulos, parrafos, tablas).",
                "requires_confirmation": True,
                "params": {
                    "path": "string (destino .docx dentro de HOME)",
                    "title": "string opcional (titulo principal del documento)",
                    "blocks": ("lista opcional de bloques: "
                               "{type:'heading',text,level(1-4)} | {type:'paragraph',text,bold?,italic?} | "
                               "{type:'table',rows:[[..],..],header?}"),
                    "content": "string opcional (alternativa simple: parrafos separados por doble salto de linea)",
                },
            },
            {
                "id": "write_xlsx",
                "description": "Crea un Excel desde filas. Una hoja simple (rows) o varias (sheets).",
                "requires_confirmation": True,
                "params": {
                    "path": "string (destino .xlsx dentro de HOME)",
                    "rows": "lista opcional de filas [[celda,..],..] para una unica hoja",
                    "sheets": "lista opcional de {name, rows:[[..],..]} para varias hojas",
                    "header": "bool opcional (pone en negrita la primera fila de cada hoja)",
                },
            },
        ]

    # ------------------------------------------------------------------
    # Lectura
    # ------------------------------------------------------------------
    async def _read_pdf(self, params: Dict[str, Any]) -> Dict[str, Any]:
        path, err = _check_readable(params.get("path"))
        if err:
            return err
        pages_spec = params.get("pages")

        def _do():
            from pypdf import PdfReader  # LAZY
            reader = PdfReader(str(path))
            total = len(reader.pages)
            indices = _parse_page_range(pages_spec, total)
            texts: list[dict] = []
            chars = 0
            truncated = False
            for i in indices:
                if len(texts) >= MAX_PDF_PAGES:
                    truncated = True
                    break
                try:
                    txt = reader.pages[i].extract_text() or ""
                except Exception as e:
                    txt = f"[error extrayendo pagina {i + 1}: {e}]"
                if chars + len(txt) > MAX_TEXT_CHARS:
                    txt = txt[: max(0, MAX_TEXT_CHARS - chars)]
                    truncated = True
                chars += len(txt)
                texts.append({"page": i + 1, "text": txt})
                if truncated:
                    break
            full_text = "\n\n".join(t["text"] for t in texts)
            has_text = bool(full_text.strip())
            return {
                "path": str(path),
                "total_pages": total,
                "pages_read": len(texts),
                "pages": texts,
                "text": full_text,
                "truncated": truncated,
                # Aviso honesto: un PDF escaneado devuelve texto vacio.
                "note": (None if has_text else
                         "El PDF no contiene texto extraible (probablemente escaneado/solo imagen). "
                         "Para leer texto de una imagen usa OCR (desktop_tool)."),
            }

        result = await asyncio.to_thread(_do)
        return {"success": True, "result": result, "error": None}

    async def _read_docx(self, params: Dict[str, Any]) -> Dict[str, Any]:
        path, err = _check_readable(params.get("path"))
        if err:
            return err

        def _do():
            import docx  # LAZY
            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables = []
            for t in doc.tables:
                rows = [[cell.text for cell in row.cells] for row in t.rows]
                tables.append(rows)

            # [S5 · NEW-1] Cabeceras y pies. Antes se omitian EN SILENCIO (a
            # diferencia de un PDF escaneado, que si lleva un `note` honesto
            # explicando por que no hay texto). En un GDD con portada, el
            # titulo del documento suele vivir justo ahi -- bastaba eso para un
            # "solo leyo una parte" sin que interviniera ningun limite de
            # tamano. Envuelto en try/except: una seccion rara de python-docx
            # no puede tumbar la lectura del cuerpo, que es lo importante.
            headers: List[str] = []
            footers: List[str] = []
            try:
                for section in doc.sections:
                    headers += [p.text for p in section.header.paragraphs if p.text.strip()]
                    footers += [p.text for p in section.footer.paragraphs if p.text.strip()]
            except Exception:
                pass

            partes = headers + paragraphs + footers
            full_text = "\n".join(partes)
            truncated = len(full_text) > MAX_TEXT_CHARS
            if truncated:
                full_text = full_text[:MAX_TEXT_CHARS]

            # Aviso honesto SIEMPRE (mismo patron que el `note` de read_pdf):
            # python-docx no expone cuadros de texto ni objetos incrustados, y
            # el que lee no tiene forma de saberlo si no se le dice.
            note = ("Extraidos cuerpo, tablas, cabeceras y pies. Los cuadros de texto y "
                    "objetos incrustados NO se extraen: si falta contenido, puede estar ahi.")
            if truncated:
                note += f" Texto recortado a {MAX_TEXT_CHARS} caracteres."

            return {
                "path": str(path),
                "paragraphs": paragraphs,
                "tables": tables,
                "headers": headers,
                "footers": footers,
                "text": full_text,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "truncated": truncated,
                "note": note,
            }

        result = await asyncio.to_thread(_do)
        return {"success": True, "result": result, "error": None}

    async def _read_xlsx(self, params: Dict[str, Any]) -> Dict[str, Any]:
        path, err = _check_readable(params.get("path"))
        if err:
            return err
        only_sheet = params.get("sheet")

        def _do():
            from openpyxl import load_workbook  # LAZY
            wb = load_workbook(str(path), read_only=True, data_only=True)
            try:
                names = wb.sheetnames
                if only_sheet is not None:
                    if only_sheet not in names:
                        raise ValueError(f"la hoja '{only_sheet}' no existe. Hojas: {names}")
                    names = [only_sheet]
                sheets = []
                for name in names:
                    ws = wb[name]
                    rows = []
                    truncated = False
                    for r in ws.iter_rows(values_only=True):
                        if len(rows) >= MAX_XLSX_ROWS:
                            truncated = True
                            break
                        # Serializable a JSON: fechas/None tal cual -> str seguro.
                        rows.append([_cell_to_jsonable(c) for c in r[:MAX_XLSX_COLS]])
                    sheets.append({"name": name, "rows": rows, "row_count": len(rows),
                                   "truncated": truncated})
                return {"path": str(path), "sheet_names": wb.sheetnames, "sheets": sheets}
            finally:
                wb.close()

        result = await asyncio.to_thread(_do)
        return {"success": True, "result": result, "error": None}

    # ------------------------------------------------------------------
    # Escritura
    # ------------------------------------------------------------------
    async def _write_docx(self, params: Dict[str, Any]) -> Dict[str, Any]:
        path, err = _check_writable(params.get("path"))
        if err:
            return err
        title = params.get("title")
        blocks = params.get("blocks")
        content = params.get("content")
        if not title and not blocks and not content:
            return {"success": False, "result": None,
                    "error": "falta contenido: da 'title', 'blocks' o 'content'"}
        if blocks is not None and not isinstance(blocks, list):
            return {"success": False, "result": None, "error": "'blocks' debe ser una lista"}
        if blocks and len(blocks) > MAX_WRITE_BLOCKS:
            return {"success": False, "result": None,
                    "error": f"demasiados bloques ({len(blocks)}, max {MAX_WRITE_BLOCKS})"}

        def _do():
            import docx  # LAZY
            doc = docx.Document()
            if title:
                doc.add_heading(str(title), level=0)
            # Camino simple: 'content' como parrafos separados por doble salto.
            if content and not blocks:
                for para in str(content).split("\n\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())
            # Camino estructurado: bloques.
            for block in (blocks or []):
                if not isinstance(block, dict):
                    continue
                btype = block.get("type", "paragraph")
                if btype == "heading":
                    lvl = block.get("level", 1)
                    try:
                        lvl = max(1, min(4, int(lvl)))
                    except (TypeError, ValueError):
                        lvl = 1
                    doc.add_heading(str(block.get("text", "")), level=lvl)
                elif btype == "table":
                    rows = block.get("rows") or []
                    if rows and isinstance(rows, list):
                        ncols = max(len(r) for r in rows if isinstance(r, list)) if rows else 0
                        if ncols:
                            tbl = doc.add_table(rows=0, cols=ncols)
                            tbl.style = "Table Grid"
                            header = block.get("header")
                            for ri, r in enumerate(rows):
                                cells = tbl.add_row().cells
                                for ci in range(ncols):
                                    val = r[ci] if isinstance(r, list) and ci < len(r) else ""
                                    cells[ci].text = "" if val is None else str(val)
                                    if header and ri == 0:
                                        for p in cells[ci].paragraphs:
                                            for run in p.runs:
                                                run.bold = True
                else:  # paragraph (default)
                    p = doc.add_paragraph()
                    run = p.add_run(str(block.get("text", "")))
                    if block.get("bold"):
                        run.bold = True
                    if block.get("italic"):
                        run.italic = True

            path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(path))
            return path.stat().st_size

        size = await asyncio.to_thread(_do)
        return {"success": True, "result": {"path": str(path), "size": size}, "error": None}

    async def _write_xlsx(self, params: Dict[str, Any]) -> Dict[str, Any]:
        path, err = _check_writable(params.get("path"))
        if err:
            return err
        rows = params.get("rows")
        sheets = params.get("sheets")
        header = bool(params.get("header"))
        if rows is None and sheets is None:
            return {"success": False, "result": None,
                    "error": "falta contenido: da 'rows' (una hoja) o 'sheets' (varias)"}

        # Normaliza a una lista de {name, rows}.
        if sheets is None:
            sheets = [{"name": "Hoja1", "rows": rows or []}]
        if not isinstance(sheets, list):
            return {"success": False, "result": None, "error": "'sheets' debe ser una lista"}
        for s in sheets:
            srows = s.get("rows") if isinstance(s, dict) else None
            if srows is not None and len(srows) > MAX_XLSX_ROWS:
                return {"success": False, "result": None,
                        "error": f"demasiadas filas ({len(srows)}, max {MAX_XLSX_ROWS})"}

        def _do():
            from openpyxl import Workbook  # LAZY
            from openpyxl.styles import Font
            wb = Workbook()
            wb.remove(wb.active)  # empezamos sin hoja por defecto
            for si, s in enumerate(sheets):
                if not isinstance(s, dict):
                    continue
                name = str(s.get("name") or f"Hoja{si + 1}")[:31]  # Excel: max 31 chars
                srows = s.get("rows") or []
                ws = wb.create_sheet(title=name)
                for r in srows[:MAX_XLSX_ROWS]:
                    ws.append([_cell_to_writable(c) for c in (r or [])[:MAX_XLSX_COLS]])
                if header and ws.max_row >= 1:
                    for cell in ws[1]:
                        cell.font = Font(bold=True)
            if not wb.sheetnames:
                wb.create_sheet(title="Hoja1")
            path.parent.mkdir(parents=True, exist_ok=True)
            wb.save(str(path))
            return path.stat().st_size

        size = await asyncio.to_thread(_do)
        return {"success": True, "result": {"path": str(path),
                                            "sheets": len(sheets), "size": size}, "error": None}


# ---------------------------------------------------------------------------
# Helpers puros (testeables sin las librerias)
# ---------------------------------------------------------------------------
def _parse_page_range(spec: Optional[str], total: int) -> List[int]:
    """'1-5'/'3'/None -> lista de indices 0-indexados dentro de [0, total)."""
    if not spec:
        return list(range(total))
    spec = str(spec).strip()
    try:
        if "-" in spec:
            a, b = spec.split("-", 1)
            start = max(1, int(a.strip()))
            end = min(total, int(b.strip()))
        else:
            start = end = int(spec)
    except (ValueError, TypeError):
        return list(range(total))
    start = max(1, start)
    end = min(total, end)
    if start > end:
        return []
    return list(range(start - 1, end))


def _cell_to_jsonable(value):
    """Valor de celda leido -> algo serializable a JSON de forma segura."""
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    return str(value)  # fechas, decimales, etc.


def _cell_to_writable(value):
    """Valor entrante -> algo que openpyxl acepte en una celda."""
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    return str(value)
